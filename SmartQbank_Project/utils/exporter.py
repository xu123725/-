from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas


def _option_label(index: int) -> str:
    return chr(ord("A") + index)


def _question_lines(questions: list[dict[str, Any]]) -> list[str]:
    lines: list[str] = []
    for idx, q in enumerate(questions, start=1):
        lines.append(f"{idx}. {q.get('stem', '')}")
        options = q.get("options", [])
        for opt_idx, opt in enumerate(options):
            lines.append(f"   {_option_label(opt_idx)}. {opt}")
        answer = q.get("answer", "")
        lines.append(f"   答案：{answer}")
        lines.append("")
    return lines


def export_paper_to_docx(questions: list[dict[str, Any]], title: str = "智能组卷试卷") -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    for idx, q in enumerate(questions, start=1):
        document.add_paragraph(f"{idx}. {q.get('stem', '')}")
        options = q.get("options", [])
        for opt_idx, opt in enumerate(options):
            document.add_paragraph(f"{_option_label(opt_idx)}. {opt}")
        document.add_paragraph(f"答案：{q.get('answer', '')}")
        document.add_paragraph("")
    bio = BytesIO()
    document.save(bio)
    return bio.getvalue()


def export_paper_to_pdf(questions: list[dict[str, Any]], title: str = "智能组卷试卷") -> bytes:
    bio = BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(bio, pagesize=A4)
    c.setFont("STSong-Light", 12)
    width, height = A4
    x = 42
    y = height - 42
    line_height = 18
    c.drawString(x, y, title)
    y -= line_height * 1.5
    for line in _question_lines(questions):
        if y < 42:
            c.showPage()
            c.setFont("STSong-Light", 12)
            y = height - 42
        c.drawString(x, y, line[:120])
        y -= line_height
    c.save()
    return bio.getvalue()
